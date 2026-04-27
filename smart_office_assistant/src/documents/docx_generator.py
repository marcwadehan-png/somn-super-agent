"""
__all__ = [
    'add_bullet_list',
    'add_footer',
    'add_header',
    'add_heading',
    'add_image',
    'add_numbered_list',
    'add_page_break',
    'add_paragraph',
    'add_table',
    'add_toc',
    'create_meeting_minutes',
    'create_project_proposal',
    'create_report',
    'save',
]

Word 文档generate器 - DOCX Generator
基于 python-docx 的文档generate功能
"""

from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from loguru import logger

@dataclass
class DocumentSection:
    """文档章节"""
    title: str
    content: str
    level: int = 1  # 标题级别 1-6
    style: Optional[str] = None

@dataclass
class TableData:
    """表格数据"""
    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str] = None

@dataclass
class ImageData:
    """图片数据"""
    path: str
    caption: Optional[str] = None
    width: Optional[float] = None  # 英寸

class DOCXGenerator:
    """
    Word 文档generate器
    
    功能:
    - 创建和编辑 Word 文档
    - 支持标题,段落,列表,表格
    - 插入图片和图表
    - 设置样式和格式
    - generate目录
    """
    
    def __init__(self, template_path: Optional[str] = None):
        """
        initgenerate器
        
        Args:
            template_path: 可选的模板文件路径
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
            logger.info(f"加载模板: {template_path}")
        else:
            self.doc = Document()
            self._setup_default_styles()
        
        self.template_path = template_path
    
    def _setup_default_styles(self):
        """设置默认样式"""
        # 设置中文字体
        styles = self.doc.styles
        
        # 正文样式
        style = styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        style.font.size = Pt(12)
        
        # 标题样式
        for i in range(1, 7):
            style_name = f'Heading {i}'
            if style_name in styles:
                style = styles[style_name]
                style.font.name = 'Microsoft YaHei'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                style.font.bold = True
                style.font.color.rgb = RGBColor(0, 0, 0)
    
    def add_heading(self, text: str, level: int = 1) -> Any:
        """
        添加标题
        
        Args:
            text: 标题文本
            level: 标题级别 1-6
        """
        heading = self.doc.add_heading(text, level=level)
        # 确保中文字体
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        return heading
    
    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        alignment: Optional[str] = None,
        bold: bool = False,
        italic: bool = False,
        font_size: Optional[int] = None
    ) -> Any:
        """
        添加段落
        
        Args:
            text: 段落文本
            style: 段落样式
            alignment: 对齐方式 (left, center, right, justify)
            bold: 是否粗体
            italic: 是否斜体
            font_size: 字体大小
        """
        p = self.doc.add_paragraph(text, style=style)
        
        # 设置对齐
        if alignment:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            p.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # 设置格式
        for run in p.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if font_size:
                run.font.size = Pt(font_size)
        
        return p
    
    def add_bullet_list(self, items: List[str], level: int = 0) -> List[Any]:
        """
        添加项目符号列表
        
        Args:
            items: 列表项
            level: 缩进级别
        """
        paragraphs = []
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            paragraphs.append(p)
        
        return paragraphs
    
    def add_numbered_list(self, items: List[str], level: int = 0) -> List[Any]:
        """
        添加编号列表
        
        Args:
            items: 列表项
            level: 缩进级别
        """
        paragraphs = []
        for item in items:
            p = self.doc.add_paragraph(item, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            paragraphs.append(p)
        
        return paragraphs
    
    def add_table(
        self,
        data: TableData,
        style: str = 'Light Grid Accent 1'
    ) -> Any:
        """
        添加表格
        
        Args:
            data: 表格数据
            style: 表格样式
        """
        table = self.doc.add_table(rows=1, cols=len(data.headers))
        table.style = style
        
        # 表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(data.headers):
            hdr_cells[i].text = header
            # 设置表头格式
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 数据行
        for row_data in data.rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                # 设置字体
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加表注
        if data.caption:
            self.add_paragraph(data.caption, alignment='center', italic=True)
        
        return table
    
    def add_image(
        self,
        image_data: ImageData,
        alignment: str = 'center'
    ) -> Any:
        """
        添加图片
        
        Args:
            image_data: 图片数据
            alignment: 对齐方式
        """
        if not Path(image_data.path).exists():
            logger.warning(f"图片不存在: {image_data.path}")
            return None
        
        # 添加图片
        if image_data.width:
            run = self.doc.add_paragraph().add_run()
            run.add_picture(image_data.path, width=Inches(image_data.width))
        else:
            run = self.doc.add_paragraph().add_run()
            run.add_picture(image_data.path)
        
        # 设置对齐
        paragraph = run._element.getparent()
        p = paragraph.getparent()
        
        # 添加图注
        if image_data.caption:
            self.add_paragraph(image_data.caption, alignment='center', italic=True)
        
        return run
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def add_toc(self, title: str = "目录"):
        """
        添加目录
        
        Note: 需要在文档generate后手动更新
        """
        self.add_heading(title, level=1)
        
        # 创建目录字段
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # 添加 TOC 字段
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        self.add_page_break()
    
    def add_header(self, text: str):
        """添加页眉"""
        section = self.doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = text
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in paragraph.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
    
    def add_footer(self, text: str = None, include_page_number: bool = True):
        """添加页脚"""
        section = self.doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        
        if include_page_number:
            # 添加页码
            run = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            instrText = OxmlElement('w:instrText')
            instrText.text = "PAGE"
            
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            
            if text:
                paragraph.add_run(f" {text}")
        else:
            paragraph.text = text or ""
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_report(
        self,
        title: str,
        sections: List[DocumentSection],
        author: str = "",
        date: Optional[str] = None,
        add_toc: bool = True
    ):
        """
        创建完整报告
        
        Args:
            title: 报告标题
            sections: 章节列表
            author: 作者
            date: 日期
            add_toc: 是否添加目录
        """
        # 封面
        self._create_cover_page(title, author, date)
        self.add_page_break()
        
        # 目录
        if add_toc:
            self.add_toc()
        
        # 正文
        for section in sections:
            self.add_heading(section.title, level=section.level)
            
            # 处理内容(支持简单换行)
            paragraphs = section.content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    self.add_paragraph(para.strip())
    
    def _create_cover_page(
        self,
        title: str,
        author: str = "",
        date: Optional[str] = None
    ):
        """创建封面"""
        # 空行
        for _ in range(6):
            self.doc.add_paragraph()
        
        # 标题
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = 'Microsoft YaHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 空行
        for _ in range(4):
            self.doc.add_paragraph()
        
        # 作者
        if author:
            author_para = self.doc.add_paragraph()
            author_run = author_para.add_run(f"作者: {author}")
            author_run.font.size = Pt(14)
            author_run.font.name = 'Microsoft YaHei'
            author_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 日期
        if date:
            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(date)
            date_run.font.size = Pt(12)
            date_run.font.name = 'Microsoft YaHei'
            date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            date_run.font.color.rgb = RGBColor(128, 128, 128)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def save(self, output_path: str) -> str:
        """
        保存文档
        
        Args:
            output_path: 输出路径
        
        Returns:
            保存的文件路径
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(output_path)
        logger.info(f"文档已保存: {output_path}")
        
        return str(output_path)

# 预设模板
class DocumentTemplates:
    """文档模板"""
    
    @staticmethod
    def create_meeting_minutes(
        meeting_title: str,
        date: str,
        attendees: List[str],
        agenda: List[str],
        discussions: List[Dict[str, str]],
        decisions: List[str],
        action_items: List[Dict[str, str]]
    ) -> DOCXGenerator:
        """创建会议纪要模板"""
        gen = DOCXGenerator()
        
        # 标题
        gen.add_heading(f"会议纪要: {meeting_title}", level=1)
        
        # 基本信息
        gen.add_paragraph(f"日期: {date}")
        gen.add_paragraph(f"参会人员: {', '.join(attendees)}")
        
        # 议程
        gen.add_heading("议程", level=2)
        gen.add_numbered_list(agenda)
        
        # 讨论内容
        gen.add_heading("讨论内容", level=2)
        for discussion in discussions:
            gen.add_heading(discussion.get('topic', ''), level=3)
            gen.add_paragraph(discussion.get('content', ''))
        
        # 决议
        gen.add_heading("决议事项", level=2)
        gen.add_numbered_list(decisions)
        
        # action项
        gen.add_heading("action项", level=2)
        action_table = TableData(
            headers=["事项", "负责人", "截止日期"],
            rows=[[
                item.get('task', ''),
                item.get('owner', ''),
                item.get('deadline', '')
            ] for item in action_items]
        )
        gen.add_table(action_table)
        
        return gen
    
    @staticmethod
    def create_project_proposal(
        project_name: str,
        background: str,
        objectives: List[str],
        scope: str,
        timeline: List[Dict[str, str]],
        budget: str
    ) -> DOCXGenerator:
        """创建项目提案模板"""
        gen = DOCXGenerator()
        
        gen.add_heading(f"项目提案: {project_name}", level=1)
        
        gen.add_heading("项目背景", level=2)
        gen.add_paragraph(background)
        
        gen.add_heading("项目目标", level=2)
        gen.add_bullet_list(objectives)
        
        gen.add_heading("项目范围", level=2)
        gen.add_paragraph(scope)
        
        gen.add_heading("项目计划", level=2)
        timeline_table = TableData(
            headers=["阶段", "开始日期", "结束日期", "主要交付物"],
            rows=[[
                item.get('phase', ''),
                item.get('start', ''),
                item.get('end', ''),
                item.get('deliverable', '')
            ] for item in timeline]
        )
        gen.add_table(timeline_table)
        
        gen.add_heading("预算估算", level=2)
        gen.add_paragraph(budget)
        
        return gen
