# -*- coding: utf-8 -*-
"""主窗口 - 导出模块

__all__ = [
    'export_document',
    'export_to_docx',
    'export_to_file',
    'export_to_html',
    'export_to_md',
    'export_to_pdf',
    'export_to_txt',
]

处理文档导出功能（txt, docx, pdf, md, html）
"""

from PySide6.QtWidgets import QDialog, QVBoxLayout, QFormLayout, QLineEdit, QComboBox, QDialogButtonBox, QMessageBox
from PySide6.QtCore import Qt

def export_document(main_window):
    """显示导出对话框并导出文档"""
    # 创建导出对话框
    dialog = QDialog(main_window)
    dialog.setWindowTitle("导出文档")
    dialog.setMinimumWidth(400)

    layout = QVBoxLayout(dialog)

    # 表单布局
    form_layout = QFormLayout()

    # 文件名输入
    filename_input = QLineEdit()
    filename_input.setText(main_window.doc_title_input.text() or "未命名")
    form_layout.addRow("文件名:", filename_input)

    # 格式选择
    format_combo = QComboBox()
    format_combo.addItems([
        "文本文件 (*.txt)",
        "Word文档 (*.docx)",
        "PDF文档 (*.pdf)",
        "Markdown (*.md)",
        "HTML (*.html)"
    ])
    form_layout.addRow("格式:", format_combo)

    layout.addLayout(form_layout)

    # 按钮
    button_box = QDialogButtonBox(
        QDialogButtonBox.Ok | QDialogButtonBox.Cancel
    )
    button_box.accepted.connect(dialog.accept)
    button_box.rejected.connect(dialog.reject)
    layout.addWidget(button_box)

    if dialog.exec() == QDialog.Accepted:
        filename = filename_input.text().strip()
        format_text = format_combo.currentText()

        # 确定文件扩展名和过滤器
        if "txt" in format_text:
            ext = ".txt"
            filter_str = "文本文件 (*.txt)"
        elif "docx" in format_text:
            ext = ".docx"
            filter_str = "Word文档 (*.docx)"
        elif "pdf" in format_text:
            ext = ".pdf"
            filter_str = "PDF文档 (*.pdf)"
        elif "md" in format_text:
            ext = ".md"
            filter_str = "Markdown (*.md)"
        elif "html" in format_text:
            ext = ".html"
            filter_str = "HTML (*.html)"
        else:
            ext = ".txt"
            filter_str = "文本文件 (*.txt)"

        # 确保文件名有扩展名
        if not filename.endswith(ext):
            filename += ext

        # 选择保存路径
        from PySide6.QtWidgets import QFileDialog
        from pathlib import Path
        file_path, _ = QFileDialog.getSaveFileName(
            main_window,
            "导出文档",
            str(main_window.current_workspace / filename),
            filter_str
        )

        if file_path:
            export_to_file(main_window, file_path, ext)

def export_to_file(main_window, file_path: str, ext: str):
    """
    导出内容到文件

    Args:
        main_window: MainWindow 实例
        file_path: 保存路径
        ext: 文件扩展名
    """
    content = main_window.editor.toPlainText()
    title = main_window.doc_title_input.text() or "未命名"

    try:
        if ext == ".txt":
            export_to_txt(file_path, content)
        elif ext == ".docx":
            export_to_docx(file_path, content, title)
        elif ext == ".pdf":
            export_to_pdf(file_path, content, title)
        elif ext == ".md":
            export_to_md(file_path, content, title)
        elif ext == ".html":
            export_to_html(file_path, content, title)

        main_window.status_bar.showMessage(f"已导出: {file_path}")
        QMessageBox.information(main_window, "导出成功", f"文档已导出到:\n{file_path}")

    except Exception as e:
        from loguru import logger
        logger.error(f"导出失败: {e}")
        QMessageBox.warning(main_window, "导出失败", f"无法导出文档: {e}")

def export_to_txt(file_path: str, content: str):
    """导出为文本文件"""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(content)

def export_to_docx(file_path: str, content: str, title: str):
    """导出为Word文档"""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # 添加标题
        doc.add_heading(title, level=0)

        # 添加内容
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        doc.save(file_path)

    except ImportError:
        # 未安装python-docx,导出为txt并提示
        txt_path = file_path.replace('.docx', '.txt')
        export_to_txt(txt_path, content)
        raise ImportError("未安装python-docx库,已导出为文本文件")

def export_to_pdf(file_path: str, content: str, title: str):
    """导出为PDF文档"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import inch

        # 注册中文字体
        try:
            pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
            chinese_font = 'SimSun'
        except Exception as e:
            try:
                pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
                chinese_font = 'SimHei'
            except Exception as e2:
                from loguru import logger
                logger.warning(f"中文字体注册失败(SimSun: {e}, SimHei: {e2})，使用 Helvetica（中文可能显示异常）")
                chinese_font = 'Helvetica'

        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )

        styles = getSampleStyleSheet()

        # 创建中文样式
        title_style = ParagraphStyle(
            'ChineseTitle',
            parent=styles['Title'],
            fontName=chinese_font,
            fontSize=18,
            spaceAfter=30
        )

        body_style = ParagraphStyle(
            'ChineseBody',
            parent=styles['BodyText'],
            fontName=chinese_font,
            fontSize=11,
            spaceAfter=12,
            leading=16
        )

        story = []

        # 添加标题
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))

        # 添加内容
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                # 处理HTML特殊字符
                p = paragraph.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(p, body_style))
                story.append(Spacer(1, 0.1 * inch))

        doc.build(story)

    except ImportError:
        # 未安装reportlab,导出为txt并提示
        txt_path = file_path.replace('.pdf', '.txt')
        export_to_txt(txt_path, content)
        raise ImportError("未安装reportlab库,已导出为文本文件")

def export_to_md(file_path: str, content: str, title: str):
    """导出为Markdown文件"""
    md_content = f"# {title}\n\n{content}"
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(md_content)

def export_to_html(file_path: str, content: str, title: str):
    """导出为HTML文件"""
    # 将换行转换为HTML标签
    html_content = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    html_content = html_content.replace('\n\n', '</p>\n<p>').replace('\n', '<br>')

    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        p {{
            margin: 1em 0;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p>{html_content}</p>
</body>
</html>"""

    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html)
