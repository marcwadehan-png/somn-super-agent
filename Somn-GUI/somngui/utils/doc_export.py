# -*- coding: utf-8 -*-
"""Somn GUI - 文档导出工具

纯文件导出函数，不依赖 Qt 或 Somn 核心。
支持 txt, md, html, docx, pdf 导出格式。
"""

from __future__ import annotations

from pathlib import Path

from loguru import logger


def export_to_txt(file_path: Path | str, content: str, title: str = "") -> None:
    """导出为文本文件"""
    header = f"{title}\n{'=' * len(title)}\n\n" if title else ""
    Path(file_path).write_text(header + content, encoding='utf-8')


def export_to_md(file_path: Path | str, content: str, title: str = "") -> None:
    """导出为 Markdown 文件"""
    md = f"# {title}\n\n{content}" if title else content
    Path(file_path).write_text(md, encoding='utf-8')


def export_to_html(file_path: Path | str, content: str, title: str = "") -> None:
    """导出为 HTML 文件"""
    safe = content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    body = safe.replace('\n\n', '</p>\n<p>').replace('\n', '<br>')

    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            color: #333;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        p {{ margin: 1em 0; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p>{body}</p>
</body>
</html>"""
    Path(file_path).write_text(html, encoding='utf-8')


def export_to_docx(file_path: Path | str, content: str, title: str = "") -> None:
    """导出为 Word 文档（需 python-docx）"""
    try:
        from docx import Document
        doc = Document()
        if title:
            doc.add_heading(title, level=0)
        for para in content.split('\n\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        doc.save(str(file_path))
    except ImportError:
        # 回退到 txt
        fallback = str(file_path).replace('.docx', '.txt')
        export_to_txt(fallback, content)
        raise ImportError("未安装 python-docx，已导出为文本文件")


def export_to_pdf(file_path: Path | str, content: str, title: str = "") -> None:
    """导出为 PDF 文档（需 reportlab）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        fallback = str(file_path).replace('.pdf', '.txt')
        export_to_txt(fallback, content)
        raise ImportError("未安装 reportlab，已导出为文本文件")

    # 尝试注册中文字体
    chinese_font = 'Helvetica'
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
        chinese_font = 'SimSun'
    except Exception:
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))
            chinese_font = 'SimHei'
        except Exception:
            pass

    doc = SimpleDocTemplate(
        str(file_path), pagesize=A4,
        rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18,
    )
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'ChineseTitle', parent=styles['Title'],
        fontName=chinese_font, fontSize=18, spaceAfter=30,
    )
    body_style = ParagraphStyle(
        'ChineseBody', parent=styles['BodyText'],
        fontName=chinese_font, fontSize=11, spaceAfter=12, leading=16,
    )

    story = []
    if title:
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))

    for para in content.split('\n\n'):
        if para.strip():
            safe = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe, body_style))
            story.append(Spacer(1, 0.1 * inch))

    doc.build(story)


# 格式映射表
EXPORT_FORMATS = {
    '.txt':  ('文本文件', export_to_txt),
    '.md':   ('Markdown', export_to_md),
    '.html': ('HTML',     export_to_html),
    '.docx': ('Word文档', export_to_docx),
    '.pdf':  ('PDF文档',  export_to_pdf),
}


def export_document(
    file_path: Path | str,
    content: str,
    fmt: str = ".txt",
    title: str = "",
) -> Path:
    """通用导出入口

    Args:
        file_path: 输出路径
        content: 文档内容
        fmt: 文件扩展名 (.txt, .md, .html, .docx, .pdf)
        title: 文档标题

    Returns:
        实际写入的文件路径
    """
    path = Path(file_path)
    fmt = fmt.lower() if fmt.startswith('.') else f'.{fmt}'

    if fmt not in EXPORT_FORMATS:
        fmt = '.txt'

    export_fn = EXPORT_FORMATS[fmt][1]
    export_fn(path, content, title=title)
    logger.info(f"文档已导出: {path} ({fmt})")
    return path
