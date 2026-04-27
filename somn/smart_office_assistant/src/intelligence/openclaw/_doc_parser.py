# -*- coding: utf-8 -*-
"""
文档解析器 - 多格式文档结构化提取

功能:
- Markdown 解析（标题/段落/列表/表格/代码块/链接）
- PDF 文本提取
- Word (.docx) 解析
- 纯文本 (.txt) 智能分段
- CSV/Excel 结构化读取
- HTML 文档解析
- 自动格式检测
- 分块与元数据提取

版本: v1.0.0
创建: 2026-04-23
"""

from __future__ import annotations
import re
import csv
import io
import json
import logging
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum

logger = logging.getLogger(__name__)


class DocFormat(Enum):
    """文档格式"""
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    HTML = "html"
    JSON = "json"
    UNKNOWN = "unknown"


@dataclass
class ParseConfig:
    """解析配置"""
    max_chunk_size: int = 2000        # 分块最大字符数
    chunk_overlap: int = 200           # 分块重叠字符数
    extract_metadata: bool = True      # 是否提取元数据
    extract_links: bool = True         # 是否提取链接
    preserve_structure: bool = True    # 是否保留结构标记
    encoding: str = "utf-8"            # 默认编码
    fallback_encodings: List[str] = field(default_factory=lambda: ["gbk", "gb2312", "latin-1"])


@dataclass
class DocChunk:
    """文档分块"""
    index: int
    content: str
    chunk_type: str = "text"           # title / paragraph / list / table / code / mixed
    heading: str = ""                  # 所属标题
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocMetadata:
    """文档元数据"""
    title: str = ""
    author: str = ""
    created: Optional[datetime] = None
    modified: Optional[datetime] = None
    word_count: int = 0
    char_count: int = 0
    section_count: int = 0
    link_count: int = 0
    format: DocFormat = DocFormat.UNKNOWN
    source_path: str = ""


@dataclass
class ParseResult:
    """解析结果"""
    raw_text: str = ""
    chunks: List[DocChunk] = field(default_factory=list)
    metadata: DocMetadata = field(default_factory=DocMetadata)
    tables: List[List[List[str]]] = field(default_factory=list)  # 提取的表格
    links: List[Dict[str, str]] = field(default_factory=list)     # 提取的链接
    error: str = ""


class DocParser:
    """文档解析器

    支持多种格式的文档解析，提供统一接口。

    示例:
        parser = DocParser()
        result = parser.parse_file("file/系统文件/paper.pdf")
        for chunk in result.chunks:
            logger.info(chunk.content)
    """

    def __init__(self, config: Optional[ParseConfig] = None):
        self.config = config or ParseConfig()
        self._parsers = {
            DocFormat.MARKDOWN: self._parse_markdown,
            DocFormat.TXT: self._parse_txt,
            DocFormat.HTML: self._parse_html,
            DocFormat.CSV: self._parse_csv,
            DocFormat.JSON: self._parse_json,
            DocFormat.PDF: self._parse_pdf,
            DocFormat.DOCX: self._parse_docx,
        }

    def parse_file(self, file_path: str) -> ParseResult:
        """解析文件

        Args:
            file_path: 文件路径

        Returns:
            ParseResult 解析结果
        """
        path = Path(file_path)
        if not path.exists():
            return ParseResult(error=f"File not found: {file_path}")

        fmt = self.detect_format(path)
        raw = self._read_file(path, fmt)

        if raw is None:
            return ParseResult(error=f"Failed to read file: {file_path}")

        # ★ v1.2 修复: 确保传入 parse_content 的内容为字符串
        if not isinstance(raw, str):
            return ParseResult(error=f"File content is not text: {type(raw).__name__}")

        return self.parse_content(raw, fmt, source_path=str(path))

    def parse_content(
        self,
        content: str,
        fmt: DocFormat,
        source_path: str = "",
    ) -> ParseResult:
        """解析文本内容

        Args:
            content: 文本内容
            fmt: 文档格式
            source_path: 来源路径（用于元数据）

        Returns:
            ParseResult
        """
        parser_fn = self._parsers.get(fmt)
        if not parser_fn:
            return ParseResult(raw_text=content, error=f"Unsupported format: {fmt}")

        try:
            return parser_fn(content, source_path)
        except Exception as e:
            # 降级为纯文本
            return ParseResult(raw_text=content, error=f"Parse error ({fmt}): {e}")

    def detect_format(self, path: Path) -> DocFormat:
        """自动检测文件格式"""
        ext_map = {
            ".md": DocFormat.MARKDOWN,
            ".markdown": DocFormat.MARKDOWN,
            ".pdf": DocFormat.PDF,
            ".docx": DocFormat.DOCX,
            ".txt": DocFormat.TXT,
            ".csv": DocFormat.CSV,
            ".tsv": DocFormat.CSV,
            ".html": DocFormat.HTML,
            ".htm": DocFormat.HTML,
            ".json": DocFormat.JSON,
        }
        return ext_map.get(path.suffix.lower(), DocFormat.UNKNOWN)

    # ─────────────────────────────────────────
    # 文件读取
    # ─────────────────────────────────────────

    def _read_file(self, path: Path, fmt: DocFormat) -> Optional[str]:
        """读取文件内容"""
        if fmt == DocFormat.PDF:
            return self._read_pdf_binary(path)
        elif fmt == DocFormat.DOCX:
            return self._read_docx_binary(path)
        else:
            return self._read_text_file(path)

    def _read_text_file(self, path: Path) -> Optional[str]:
        """读取文本文件（自动编码检测）"""
        encodings = [self.config.encoding] + self.config.fallback_encodings
        for enc in encodings:
            try:
                return path.read_text(encoding=enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return None

    def _read_pdf_binary(self, path: Path) -> Optional[str]:
        """读取PDF文件"""
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(str(path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            # PyPDF2不可用，尝试pdfplumber
            try:
                import pdfplumber
                pages = []
                with pdfplumber.open(str(path)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            pages.append(text)
                return "\n\n".join(pages)
            except ImportError:
                return None
        except Exception:
            return None

    def _read_docx_binary(self, path: Path) -> Optional[str]:
        """读取DOCX文件"""
        try:
            from docx import Document
            doc = Document(str(path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # 提取表格
            tables_text = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    rows.append(cells)
                tables_text.append(rows)

            # 拼接
            content = "\n\n".join(paragraphs)
            if tables_text:
                content += "\n\n" + "\n\n".join(
                    "\n".join("\t".join(cell for cell in row) for row in table)
                    for table in tables_text
                )
            return content
        except ImportError:
            return None
        except Exception:
            return None

    # ─────────────────────────────────────────
    # 格式解析器
    # ─────────────────────────────────────────

    def _parse_markdown(self, content: str, source_path: str = "") -> ParseResult:
        """解析Markdown文档"""
        result = ParseResult(raw_text=content)
        result.metadata.format = DocFormat.MARKDOWN
        result.metadata.source_path = source_path

        chunks = []
        current_heading = ""
        current_content = []
        current_type = "paragraph"
        tables = []
        links = []

        lines = content.split("\n")
        i = 0

        while i < len(lines):
            line = lines[i]

            # 标题
            heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
            if heading_match:
                # 保存之前的内容
                if current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        chunks.append(DocChunk(
                            index=len(chunks),
                            content=text,
                            chunk_type=current_type,
                            heading=current_heading,
                        ))
                    current_content = []

                level = len(heading_match.group(1))
                current_heading = heading_match.group(2).strip()
                chunks.append(DocChunk(
                    index=len(chunks),
                    content=current_heading,
                    chunk_type="title",
                    heading=current_heading,
                    metadata={"level": level},
                ))
                i += 1
                continue

            # 代码块
            if line.strip().startswith("```"):
                lang = line.strip()[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                code = "\n".join(code_lines)
                chunks.append(DocChunk(
                    index=len(chunks),
                    content=code,
                    chunk_type="code",
                    heading=current_heading,
                    metadata={"language": lang},
                ))
                i += 1
                continue

            # 表格
            if "|" in line and line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                table = self._parse_markdown_table(table_lines)
                if table:
                    tables.append(table)
                    chunks.append(DocChunk(
                        index=len(chunks),
                        content=self._table_to_text(table),
                        chunk_type="table",
                        heading=current_heading,
                    ))
                continue

            # 链接提取
            link_matches = re.findall(r"\[([^\]]+)\]\(([^)]+)\)", line)
            for text, url in link_matches:
                links.append({"text": text, "url": url})

            # 列表
            if re.match(r"^(\s*)([-*+]|\d+\.)\s+", line):
                if current_type != "list" and current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        chunks.append(DocChunk(
                            index=len(chunks),
                            content=text,
                            chunk_type=current_type,
                            heading=current_heading,
                        ))
                    current_content = []
                current_type = "list"
                current_content.append(line.strip())
                i += 1
                continue

            # 普通段落
            if line.strip():
                if current_type == "list" and current_content:
                    text = "\n".join(current_content).strip()
                    if text:
                        chunks.append(DocChunk(
                            index=len(chunks),
                            content=text,
                            chunk_type="list",
                            heading=current_heading,
                        ))
                    current_content = []
                current_type = "paragraph"
                current_content.append(line.strip())
            elif current_content:
                text = "\n".join(current_content).strip()
                if text:
                    chunks.append(DocChunk(
                        index=len(chunks),
                        content=text,
                        chunk_type=current_type,
                        heading=current_heading,
                    ))
                current_content = []
                current_type = "paragraph"

            i += 1

        # 保存最后一个块
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                chunks.append(DocChunk(
                    index=len(chunks),
                    content=text,
                    chunk_type=current_type,
                    heading=current_heading,
                ))

        # 分块
        result.chunks = self._chunk_list(chunks) if len(chunks) > 1 else chunks
        result.tables = tables
        result.links = links
        result.metadata.word_count = len(content.split())
        result.metadata.char_count = len(content)
        result.metadata.section_count = sum(1 for c in result.chunks if c.chunk_type == "title")
        result.metadata.link_count = len(links)

        # 提取标题
        if chunks and chunks[0].chunk_type == "title":
            result.metadata.title = chunks[0].content

        return result

    def _parse_txt(self, content: str, source_path: str = "") -> ParseResult:
        """解析纯文本"""
        result = ParseResult(raw_text=content)
        result.metadata.format = DocFormat.TXT
        result.metadata.source_path = source_path

        # 智能分段：按空行或连续缩进
        paragraphs = re.split(r"\n\s*\n|\n{2,}", content)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        chunks = []
        for i, para in enumerate(paragraphs):
            chunks.append(DocChunk(
                index=i,
                content=para,
                chunk_type="paragraph",
            ))

        result.chunks = self._chunk_list(chunks)
        result.metadata.word_count = len(content.split())
        result.metadata.char_count = len(content)
        result.metadata.title = Path(source_path).stem if source_path else ""

        return result

    def _parse_html(self, content: str, source_path: str = "") -> ParseResult:
        """解析HTML文档"""
        result = ParseResult(raw_text=content)
        result.metadata.format = DocFormat.HTML
        result.metadata.source_path = source_path

        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "html.parser")

            # 提取标题
            title = ""
            if soup.title and soup.title.string:
                title = soup.title.string.strip()

            # 提取正文
            body = soup.find("body") or soup
            # 移除噪声
            for tag in body.find_all(["script", "style", "nav", "footer", "header", "aside"]):
                tag.decompose()

            text = body.get_text(separator="\n", strip=True)

            # 智能分段
            paragraphs = [p.strip() for p in text.split("\n") if p.strip() and len(p.strip()) > 10]

            chunks = []
            for i, para in enumerate(paragraphs):
                chunks.append(DocChunk(
                    index=i,
                    content=para,
                    chunk_type="paragraph",
                ))

            # 提取链接
            links = []
            for a in soup.find_all("a", href=True):
                links.append({"text": a.get_text(strip=True), "url": a["href"]})

            result.chunks = self._chunk_list(chunks)
            result.links = links
            result.metadata.title = title
            result.metadata.word_count = len(text.split())
            result.metadata.char_count = len(text)
            result.metadata.link_count = len(links)

        except ImportError:
            result.chunks = [DocChunk(index=0, content=content[:self.config.max_chunk_size], chunk_type="raw")]

        return result

    def _parse_csv(self, content: str, source_path: str = "") -> ParseResult:
        """解析CSV/TSV文件"""
        result = ParseResult(raw_text=content)
        result.metadata.format = DocFormat.CSV
        result.metadata.source_path = source_path

        try:
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
            if not rows:
                return result

            result.tables = [rows]
            result.chunks = [DocChunk(
                index=0,
                content=self._table_to_text(rows),
                chunk_type="table",
                metadata={"rows": len(rows), "columns": len(rows[0]) if rows else 0},
            )]
            result.metadata.word_count = sum(len(r) for r in rows)
            result.metadata.title = Path(source_path).stem if source_path else "CSV Data"
        except Exception as e:
            result.error = f"CSV parse error: {e}"

        return result

    def _parse_json(self, content: str, source_path: str = "") -> ParseResult:
        """解析JSON文件"""
        result = ParseResult(raw_text=content)
        result.metadata.format = DocFormat.JSON
        result.metadata.source_path = source_path

        try:
            data = json.loads(content)
            # 递归提取所有文本值
            texts = self._extract_json_texts(data)
            combined = "\n".join(texts)

            result.chunks = self._chunk_text(combined)
            result.metadata.word_count = len(combined.split())
            result.metadata.char_count = len(combined)
        except json.JSONDecodeError as e:
            result.error = f"JSON parse error: {e}"

        return result

    def _parse_pdf(self, content: str, source_path: str = "") -> ParseResult:
        """解析PDF文本（已由_read_pdf_binary提取为纯文本）"""
        return self._parse_txt(content, source_path)

    def _parse_docx(self, content: str, source_path: str = "") -> ParseResult:
        """解析DOCX文本（已由_read_docx_binary提取为纯文本）"""
        return self._parse_txt(content, source_path)

    # ─────────────────────────────────────────
    # 分块与辅助
    # ─────────────────────────────────────────

    def _chunk_list(self, chunks: List[DocChunk]) -> List[DocChunk]:
        """对已有分块列表进行合并/拆分以适应大小限制"""
        result = []
        buffer = ""
        buffer_type = ""
        buffer_heading = ""
        buffer_meta = {}
        idx = 0

        for chunk in chunks:
            # 标题和代码块保持独立
            if chunk.chunk_type in ("title", "code", "table"):
                if buffer:
                    result.append(DocChunk(
                        index=idx, content=buffer.strip(),
                        chunk_type=buffer_type, heading=buffer_heading,
                    ))
                    idx += 1
                    buffer = ""
                result.append(chunk)
                idx += 1
                continue

            text = chunk.content
            if len(buffer) + len(text) > self.config.max_chunk_size and buffer:
                result.append(DocChunk(
                    index=idx, content=buffer.strip(),
                    chunk_type=buffer_type, heading=buffer_heading,
                ))
                idx += 1
                # 保留重叠部分
                if self.config.chunk_overlap > 0:
                    buffer = buffer[-self.config.chunk_overlap:] + "\n" + text
                else:
                    buffer = text
            else:
                buffer = buffer + "\n" + text if buffer else text
                buffer_type = chunk.chunk_type
                buffer_heading = chunk.heading or buffer_heading
                buffer_meta = chunk.metadata

        if buffer.strip():
            result.append(DocChunk(
                index=idx, content=buffer.strip(),
                chunk_type=buffer_type, heading=buffer_heading,
            ))

        return result

    def _chunk_text(self, text: str) -> List[DocChunk]:
        """将纯文本分块"""
        size = self.config.max_chunk_size
        overlap = self.config.chunk_overlap
        chunks = []

        if len(text) <= size:
            chunks.append(DocChunk(index=0, content=text, chunk_type="paragraph"))
            return chunks

        start = 0
        idx = 0
        while start < len(text):
            end = min(start + size, len(text))
            # 尝试在句号/换行处断开
            if end < len(text):
                for sep in ["\n", "。", ".", "！", "！", "？", "?"]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start:
                        end = last_sep + 1
                        break

            chunks.append(DocChunk(
                index=idx,
                content=text[start:end].strip(),
                chunk_type="paragraph",
            ))
            idx += 1
            start = end - overlap if overlap > 0 else end

        return chunks

    def _parse_markdown_table(self, lines: List[str]) -> Optional[List[List[str]]]:
        """解析Markdown表格"""
        if len(lines) < 2:
            return None

        rows = []
        for line in lines:
            # 跳过分隔行
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if cells:
                rows.append(cells)

        return rows if len(rows) >= 1 else None

    @staticmethod
    def _table_to_text(table: List[List[str]]) -> str:
        """将表格转为文本"""
        return "\n".join("\t".join(cell for cell in row) for row in table)

    @staticmethod
    def _extract_json_texts(data: Any, prefix: str = "") -> List[str]:
        """递归提取JSON中的所有文本值"""
        texts = []
        if isinstance(data, dict):
            for k, v in data.items():
                key_path = f"{prefix}.{k}" if prefix else k
                if isinstance(v, str):
                    texts.append(f"{key_path}: {v}")
                elif isinstance(v, (dict, list)):
                    texts.extend(DocParser._extract_json_texts(v, key_path))
        elif isinstance(data, list):
            for i, item in enumerate(data):
                texts.extend(DocParser._extract_json_texts(item, f"{prefix}[{i}]"))
        return texts


__all__ = [
    'DocParser',
    'ParseConfig',
    'ParseResult',
    'DocChunk',
    'DocMetadata',
    'DocFormat',
]
